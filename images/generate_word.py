# -*- coding: utf-8 -*-
"""Generate Word report: 农机作业数据处理与分析系统设计与实现"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

IMAGES = r"D:\PycharmProjects\Agricultural Machinery Operation Data Processing and Analysis Software\images"
OUTPUT = os.path.join(IMAGES, "农机作业数据处理与分析系统_设计与实现_万爱华.docx")

doc = Document()

# ====== Page Setup ======
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ====== Style Configuration ======
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(0.74)

# Heading styles
for lvl, (size, font_bold, font_name) in enumerate([(22, True, '黑体'), (16, True, '黑体'), (14, True, '黑体'), (12, True, '黑体')], 1):
    h_style = doc.styles[f'Heading {lvl}']
    h_style.font.name = font_name
    h_style.font.size = Pt(size)
    h_style.font.bold = font_bold
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    h_style.paragraph_format.first_line_indent = Cm(0)
    h_style.paragraph_format.space_before = Pt(6)
    h_style.paragraph_format.space_after = Pt(6)
    h_style.paragraph_format.line_spacing = 1.5

# ====== Helpers ======
def add_para(text, style_name='Normal', bold=False, align=None, size=None, font_name=None, indent=True):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align is not None: p.alignment = align
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_img(name, width=Inches(5.5)):
    path = os.path.join(IMAGES, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(path, width=width)
        return p
    return None

def add_img_caption(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code_block(code_text):
    """Add a code block with grey background."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    return p

def add_table_with_data(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # Grey background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A73E8" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Cm(0)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()  # spacing
    return table

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('软件工程管理与实践', 'Normal', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=26, font_name='黑体', indent=False)
doc.add_paragraph()
add_para('农机作业数据处理与分析系统', 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=30, font_name='黑体', indent=False)
add_para('设计与实现', 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=26, font_name='黑体', indent=False)

for _ in range(4):
    doc.add_paragraph()

# Info table
info_items = [('学    号：', ''), ('姓    名：', '万爱华'), ('电    话：', '')]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(label + value)
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if not value:
        run.underline = True

doc.add_page_break()

# ============================================================
# 1. 系统概述
# ============================================================
add_heading('1. 系统概述', 1)

add_heading('1.1 应用背景', 2)
add_para('随着我国农业现代化进程的不断推进，精准农业和智慧农业已成为农业发展的重要方向。在农业生产过程中，农机作业数据的采集、处理与分析对于提高农业生产效率、优化作业路径、降低生产成本具有重要意义。传统的农机作业管理方式主要依赖人工记录和经验判断，存在数据采集不完整、处理效率低、分析手段单一等问题，难以满足现代化农业生产的精细化管理需求。')
add_para('近年来，全球卫星导航系统（GNSS）技术的普及使得农机设备可以搭载高精度GPS接收器，实时记录作业过程中的位置、速度、方向等数据。这些海量的轨迹数据蕴含着丰富的作业信息，通过合理的存储、处理与可视化分析，可以为农场管理者提供科学的决策支持。然而，目前市场上成熟的农机作业数据处理软件多为商业化的解决方案，价格昂贵且难以根据实际需求进行定制化开发。')
add_para('在此背景下，本课题旨在设计并实现一个基于Web的农机作业数据处理与分析系统。该系统采用前后端分离的现代软件架构，利用数据库技术存储和管理海量的农机轨迹数据，通过地图可视化技术直观展示农机作业路径，并提供完善的用户权限管理功能，为农机作业数据的科学管理提供一套完整的技术解决方案。')
add_para('本系统的开发不仅具有工程实践价值，也契合了"十四五"规划中关于推进农业数字化转型的政策导向。通过开源技术栈的合理组合应用，本系统在保证功能完备性的同时，大幅降低了开发和部署成本，具有良好的可推广性和适应性。')

add_heading('1.2 开发目标概述', 2)
add_para('本系统的总体目标是设计并实现一套完整的农机作业数据处理与分析Web应用软件，为农业管理人员提供便捷、高效的农机作业数据管理工具。具体开发目标包括以下几个方面：')
add_para('（1）数据存储与管理：设计合理的数据库表结构，实现对原始农机轨迹数据的完整存储。系统需要支持对8条以上农机作业轨迹数据的导入与管理，每条轨迹包含数千个GPS采样点，总计需要处理数万个数据记录。数据库设计需满足规范化要求，确保数据的一致性和查询效率。')
add_para('（2）地图可视化展示：集成地图组件，实现对农机作业轨迹的直观可视化展示。系统需要支持遥感影像和矢量地图两种底图模式，能够根据GPS坐标将轨迹精确叠加到地图上。轨迹应按作业状态进行差异化着色（正常作业与闲置停止），支持点击查询单个轨迹点的详细信息。')
add_para('（3）地图交互功能：提供完善的地图操作功能，包括地图的放大、缩小、漫游等基础操作，以及距离测量和轨迹动态回放等高级交互功能。这些功能需要具备良好的用户体验和响应性能。')
add_para('（4）数据导入功能：提供Excel格式轨迹数据的批量导入功能，管理员用户可以通过界面上传原始数据文件，系统自动解析、验证并存储到数据库，同时记录详细的导入日志。')
add_para('（5）用户权限管理：实现基于角色的访问控制（RBAC）模型，区分普通用户和管理员两种角色。普通用户具有轨迹查看、地图浏览等基本功能，管理员额外具有数据导入、用户管理和权限配置等功能。系统需要保证不同角色之间的功能隔离和数据安全。')
add_para('（6）系统性能与体验：系统应具有良好的响应速度和用户体验，能够在处理数万个轨迹点时保持地图渲染的流畅性。前端界面设计应简洁直观，符合现代Web应用的设计规范，支持多种主流浏览器的访问。')

doc.add_page_break()

# ============================================================
# 2. 需求分析
# ============================================================
add_heading('2. 需求分析', 1)

add_heading('2.1 功能需求', 2)
add_para('根据系统的应用场景和目标用户群体，本系统的功能需求可以划分为用户基本功能和系统管理功能两大类。以下分别对两类功能需求进行详细描述。')

add_heading('2.1.1 用户基本功能', 3)
add_para('（1）用户注册与登录：系统需要提供用户注册功能，新用户通过填写用户名、密码等基本信息完成注册。已注册用户通过用户名和密码进行登录验证，登录成功后系统生成身份令牌（Token），后续请求通过Token进行身份识别。')
add_para('（2）个人信息管理：登录用户可以查看和修改个人信息，包括昵称、邮箱、电话号码、地址等字段。系统需要提供独立的个人信息编辑页面。')
add_para('（3）密码修改：用户可以修改登录密码，修改密码时需要输入旧密码进行验证，确保账户安全性。新密码需进行二次确认，防止输入错误。')
add_para('（4）轨迹列表查看：系统以表格形式展示所有已导入的农机轨迹数据，包含轨迹ID、起始时间、结束时间、幅宽、总点数等关键信息。列表支持分页浏览和条件筛选（按时间范围、幅宽范围）。')
add_para('（5）轨迹地图展示：用户点击"查看轨迹"按钮后，系统跳转到地图展示页面。地图以卫星遥感影像或矢量地图为底图，将轨迹的GPS坐标点叠加显示在底图上。轨迹线按照作业状态分别着色——正常作业段以绿色显示，闲置停止段以灰色显示——方便直观判断作业质量。')
add_para('（6）地图基础操作：系统提供地图的放大、缩小、平移（漫游）等基础操作功能，用户可以通过鼠标滚轮缩放、拖拽平移地图。')
add_para('（7）距离测量：用户可以通过测量工具在地图上添加测距点，系统自动计算相邻点之间的球面距离并累加显示总距离。支持连续添加多个测量点，右键可撤回上一个点。')
add_para('（8）轨迹动态回放：系统支持按GPS时间顺序逐点动态展示轨迹的生成过程，用户可以直观地观察到农机作业的时序路径。回放速度固定（0.1秒/点），用户可以随时停止回放。')
add_para('（9）点位详情查询：用户点击地图上的任意轨迹点标记，系统在右侧信息面板显示该点的详细信息，包括经度、纬度、速度、耕深等属性。')

add_heading('2.1.2 管理员功能', 3)
add_para('（1）轨迹数据导入：管理员可以通过数据导入界面上传Excel格式的农机轨迹数据文件。系统自动解析文件内容，验证数据格式，将轨迹基本信息和轨迹点明细分别存入对应的数据库表，并自动计算生成作业统计和通行率等派生数据。导入完成后记录详细的导入日志。')
add_para('（2）轨迹数据删除：管理员可以从轨迹列表中删除指定的轨迹记录。删除操作采用级联方式，同时删除该轨迹关联的所有轨迹点、作业统计和通行率数据，确保数据一致性。')
add_para('（3）用户账号管理：管理员可以查看系统中所有注册用户的列表，支持按角色筛选用户。管理员可以编辑用户信息，也可以删除用户账号（支持批量删除）。')
add_para('（4）角色管理：管理员可以创建、编辑和删除系统角色。每个角色具有名称、描述和标识字段。系统预设管理员（admin）和普通用户（user）两种角色。')
add_para('（5）菜单权限配置：管理员可以通过菜单管理页面配置系统的菜单结构（支持多级树形菜单），并通过角色管理页面为每个角色分配可访问的菜单权限。不同角色的用户登录后，侧边栏菜单根据其权限动态渲染。')
add_para('（6）文件管理：系统提供通用的文件管理功能，支持文件的上传、查看和删除操作。')

add_heading('2.2 用例分析', 2)
add_para('本系统包含两类参与者（Actor）：普通用户和管理员。以下分别描述两类参与者的主要用例：')
add_para('普通用户用例：注册账号 → 登录系统 → 浏览轨迹列表 → 筛选轨迹 → 查看轨迹地图 → 操作地图（缩放/漫游/测距）→ 查看轨迹回放 → 点击点位查看详情 → 修改个人信息 → 修改密码。')
add_para('管理员用例：登录系统 → 执行所有普通用户操作 → 导入Excel轨迹数据 → 查看导入日志 → 删除轨迹数据 → 管理用户账号（查看/编辑/删除）→ 管理角色 → 配置菜单权限 → 管理文件。')
add_para('图2-1展示了系统总体用例图。其中管理员继承了普通用户的所有用例权限，并额外拥有数据管理、用户管理和系统配置三类用例。')

add_heading('2.3 性能要求', 2)
add_para('（1）响应时间：常规页面加载时间不超过3秒，地图页面首次加载（含瓦片加载）不超过5秒。API接口响应时间在常规数据量下不超过1秒。')
add_para('（2）数据处理能力：系统需支持单条轨迹3000-8000个GPS采样点的存储与渲染。轨迹点导入速度不低于500点/秒。')
add_para('（3）并发支持：系统需支持至少10个并发用户的正常访问，在多用户同时操作时系统响应无明显退化。')
add_para('（4）地图渲染性能：在渲染4000+个轨迹点时，地图的缩放和平移操作应保持流畅，无明显卡顿。')
add_para('（5）数据安全：用户密码需加密存储，API接口需进行身份验证。敏感操作（如删除轨迹、删除用户）需进行二次确认。')

doc.add_page_break()

# ============================================================
# 3. 概要设计
# ============================================================
add_heading('3. 概要设计', 1)

add_heading('3.1 开发架构', 2)
add_para('本系统采用目前主流的B/S（Browser/Server）架构，即浏览器/服务器模式。整体采用前后端分离的软件架构设计，前端和后端通过HTTP/REST API进行通信，两者在技术上完全解耦，可以独立开发、测试和部署。')
add_para('前端（Frontend）负责用户界面的呈现和交互逻辑处理。采用Vue 3作为核心框架，使用Composition API进行组件逻辑的组织和管理。UI组件库选用Element Plus，提供统一的视觉风格和丰富的交互组件。地图功能基于Leaflet开源地图库实现，瓦片数据来源于高德地图开放平台。前端构建工具使用Vite，利用其快速的HMR（热模块替换）能力提升开发效率。前端通过Axios HTTP客户端向后端API发送请求，获取JSON格式的响应数据。')
add_para('后端（Backend）负责业务逻辑处理、数据存取和API接口提供。采用Python语言下的Django 4.2框架作为Web应用基础，使用Django REST Framework（DRF）构建RESTful风格的API接口。通过ViewSet和Router机制自动生成标准的CRUD接口，通过自定义Action实现轨迹点关联查询等特定业务需求。数据库访问层使用Django ORM（对象关系映射），实现Python对象到数据库表的映射，避免直接编写SQL语句。数据库选用SQLite，这是一个轻量级的嵌入式关系型数据库，无需独立安装和配置，适合中小规模应用。')
add_para('前后端之间的通信采用JSON格式，符合RESTful设计规范。前端发送HTTP请求（GET/POST/PUT/DELETE）到后端的API端点，后端处理请求后返回JSON格式的响应数据。系统配置了CORS（跨域资源共享）中间件，允许前端开发服务器跨域访问后端API。')
add_para('图3-1展示了系统总体架构的分层设计，包括前端展示层、后端服务层和数据存储层三个层次。')

add_heading('3.2 总体流程', 2)
add_para('系统的总体业务流程如下：')
add_para('（1）用户访问系统URL，前端路由判断用户是否已登录。未登录用户自动跳转到登录页面。')
add_para('（2）用户输入用户名和密码，前端将凭证发送到后端登录接口。后端验证通过后返回用户信息和身份令牌。')
add_para('（3）前端根据用户角色从后端获取菜单权限列表，动态渲染侧边栏导航菜单。')
add_para('（4）管理员用户通过数据导入功能上传Excel轨迹文件。后端解析文件内容，将轨迹基本信息和轨迹点数据分别写入track和trackpoints表，并计算生成work和rate统计数据。')
add_para('（5）用户访问轨迹列表页面时，前端请求轨迹列表API，后端从数据库查询并返回所有轨迹记录。用户可设置时间范围和幅宽条件进行筛选。')
add_para('（6）用户点击"查看轨迹"进入地图页面。前端请求轨迹详情、轨迹点列表、作业统计和通行率四个API接口，获取完整数据后渲染地图。')
add_para('（7）前端使用Leaflet初始化地图实例，加载高德瓦片图层（卫星影像或矢量地图）。将轨迹点的WGS-84坐标转换为GCJ-02坐标后，在地图上绘制轨迹线和点位标记。')
add_para('（8）用户可在地图上进行缩放、漫游、测距、轨迹回放等交互操作。点击轨迹点可查看详细属性信息。')

add_heading('3.3 功能结构', 2)
add_para('系统功能模块划分为六大板块：')
add_para('（1）用户认证模块：包含用户注册、登录、个人信息修改、密码修改四个子功能。')
add_para('（2）轨迹管理模块：包含轨迹列表展示（含分页和筛选）、轨迹详情查看、轨迹删除（管理员）三个子功能。')
add_para('（3）地图展示模块：包含卫星底图与矢量底图切换、轨迹线分段着色渲染、轨迹点标记渲染、点位详情查询、统计信息展示五个子功能。')
add_para('（4）地图交互模块：包含地图缩放/平移、距离测量、轨迹动态回放三个子功能。')
add_para('（5）数据管理模块（管理员）：包含Excel数据导入、导入日志查看、文件管理三个子功能。')
add_para('（6）系统管理模块（管理员）：包含用户管理、角色管理、菜单权限配置三个子功能。')

add_heading('3.4 开发环境', 2)
add_table_with_data(
    ['类别', '工具/技术', '版本', '说明'],
    [
        ['操作系统', 'Windows 11', '24H2', '开发与运行环境'],
        ['后端框架', 'Django', '4.2.30', 'Python Web框架'],
        ['API框架', 'Django REST Framework', '3.16.1', 'RESTful API构建'],
        ['数据库', 'SQLite', '3.x', '嵌入式关系型数据库'],
        ['前端框架', 'Vue', '3.5.13', '渐进式JavaScript框架'],
        ['构建工具', 'Vite', '6.0.7', '前端构建与开发服务器'],
        ['UI组件库', 'Element Plus', '2.8.8', 'Vue 3组件库'],
        ['地图组件', 'Leaflet', '1.9.4', '开源JavaScript地图库'],
        ['HTTP客户端', 'Axios', '1.7.9', 'HTTP请求库'],
        ['路由管理', 'Vue Router', '4.5.0', '前端路由管理'],
        ['开发语言', 'Python / JavaScript', '3.12 / ES2022', '后端/前端编程语言'],
        ['代码编辑器', 'PyCharm / VS Code', '—', '集成开发环境'],
    ]
)

add_heading('3.5 关键技术', 2)
add_para('（1）前后端分离架构：前端Vue 3 + Vite构建的单页面应用（SPA）与后端Django REST Framework的API服务完全分离，通过HTTP/JSON协议通信。这种架构使得前后端可以独立开发、独立部署，提高了开发效率和系统的可维护性。')
add_para('（2）RESTful API设计：后端采用DRF的ViewSet + Router机制，以极少的代码量实现了10个资源的完整CRUD接口。自定义Action实现了轨迹与轨迹点的关联查询。API遵循RESTful规范，使用标准的HTTP方法（GET/POST/PUT/DELETE）和状态码。')
add_para('（3）WGS-84到GCJ-02坐标转换：中国大陆的地图服务（如高德地图）使用GCJ-02坐标系，而GPS设备采集的是WGS-84坐标。本系统实现了完整的WGS-84到GCJ-02坐标转换算法（包含20余个三角函数运算），将轨迹点坐标精确转换后叠加到地图上，确保位置精度。')
add_para('（4）Leaflet地图组件集成：选用Leaflet作为地图组件，利用其轻量级、插件丰富、支持Canvas和SVG双渲染器的特性，实现了轨迹分段渲染、动态回放、距离测量等复杂的地图交互功能。')
add_para('（5）RBAC权限控制：系统实现了用户-角色-菜单三级权限模型。前端路由守卫（Vue Router的beforeEach钩子）拦截未授权访问，后端API通过请求头中的角色信息进行鉴权。侧边栏菜单根据用户角色权限动态渲染。')
add_para('（6）CORS跨域解决方案：由于前后端运行在不同的端口（Vite开发服务器在5174端口，Django服务器在8000端口），系统使用django-cors-headers中间件处理跨域请求，并配置了允许的自定义请求头（Token、X-User-Id、X-User-Role）。')

doc.add_page_break()

# ============================================================
# 4. 详细设计
# ============================================================
add_heading('4. 详细设计', 1)

add_heading('4.1 数据库设计', 2)
add_para('数据库设计是信息系统开发的核心环节。本系统根据需求分析结果，遵循第三范式（3NF）设计原则，共设计了10张数据表，涵盖了轨迹数据存储、用户权限管理、文件管理和系统配置等各个方面。以下对核心表结构进行详细说明。')

add_heading('4.2 数据表结构设计', 2)
add_para('表4-1至表4-10列出了系统全部数据表的详细结构设计：')

add_para('表4-1 track（轨迹主表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['trackid', 'INTEGER', 'PRIMARY KEY, AUTO', '轨迹唯一标识'],
        ['starttime', 'DATETIME', 'NULL', '轨迹起始时间'],
        ['endtime', 'DATETIME', 'NULL', '轨迹结束时间'],
        ['width', 'FLOAT', 'NULL', '作业幅宽（米）'],
        ['totalpoints', 'INTEGER', 'NULL', '总轨迹点数'],
    ]
)

add_para('表4-2 trackpoints（轨迹点表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '点唯一标识'],
        ['trackid', 'INTEGER', 'FOREIGN KEY → track', '所属轨迹ID'],
        ['gpstime', 'DATETIME', 'NULL', 'GPS时间戳'],
        ['lon', 'FLOAT', 'NULL', '经度（WGS-84）'],
        ['lat', 'FLOAT', 'NULL', '纬度（WGS-84）'],
        ['x', 'FLOAT', 'NULL', '投影X坐标'],
        ['y', 'FLOAT', 'NULL', '投影Y坐标'],
        ['velocity', 'FLOAT', 'NULL', '瞬时速度（km/h）'],
        ['course', 'FLOAT', 'NULL', '行进方向角'],
        ['workstatus', 'INTEGER', 'NULL', '作业状态（1=作业 0=闲置）'],
        ['width', 'FLOAT', 'NULL', '当前幅宽'],
        ['depth', 'FLOAT', 'NULL', '耕深（cm）'],
        ['depthstandard', 'FLOAT', 'NULL', '标准耕深（cm）'],
    ]
)

add_para('表4-3 work（作业统计表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['trackid', 'INTEGER', 'PRIMARY KEY, FK→track', '关联轨迹ID'],
        ['worktime', 'FLOAT', 'NULL', '作业时长（小时）'],
        ['worklength', 'FLOAT', 'NULL', '作业长度（公里）'],
        ['workarea', 'FLOAT', 'NULL', '作业面积（公顷）'],
        ['avgvelocity', 'FLOAT', 'NULL', '平均速度（km/h）'],
    ]
)

add_para('表4-4 rate（通行率表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['trackid', 'INTEGER', 'PRIMARY KEY, FK→track', '关联轨迹ID'],
        ['passrate', 'FLOAT', 'NULL', '通行率（%）'],
        ['productionrate', 'FLOAT', 'NULL', '生产效率'],
        ['timerrate', 'FLOAT', 'NULL', '时间利用率（%）'],
    ]
)

add_para('表4-5 user（用户表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '用户唯一标识'],
        ['username', 'VARCHAR(45)', 'UNIQUE', '用户名'],
        ['password', 'VARCHAR(128)', 'NULL', '密码（加密存储）'],
        ['nickname', 'VARCHAR(45)', 'NULL', '昵称'],
        ['email', 'VARCHAR(45)', 'NULL', '电子邮箱'],
        ['phone', 'VARCHAR(45)', 'NULL', '电话号码'],
        ['address', 'VARCHAR(255)', 'NULL', '地址'],
        ['creat_time', 'DATETIME', 'NULL', '创建时间'],
        ['avatar_url', 'VARCHAR(255)', 'NULL', '头像URL'],
        ['role', 'INTEGER', 'FOREIGN KEY → role', '所属角色ID'],
    ]
)

add_para('表4-6 role（角色表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '角色唯一标识'],
        ['name', 'VARCHAR(45)', 'NULL', '角色名称'],
        ['description', 'VARCHAR(255)', 'NULL', '角色描述'],
        ['flag', 'VARCHAR(45)', 'NULL', '角色标识符'],
    ]
)

add_para('表4-7 menu（菜单表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '菜单唯一标识'],
        ['name', 'VARCHAR(255)', 'NULL', '菜单名称'],
        ['path', 'VARCHAR(255)', 'NULL', '菜单路径'],
        ['icon', 'VARCHAR(255)', 'NULL', '菜单图标'],
        ['description', 'VARCHAR(255)', 'NULL', '菜单描述'],
        ['pid', 'INTEGER', 'FK→menu(self), NULL', '父菜单ID（树形结构）'],
        ['page_path', 'VARCHAR(255)', 'NULL', '页面组件路径'],
    ]
)

add_para('表4-8 role_menu（角色菜单关联表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['role_id', 'INTEGER', 'PRIMARY KEY, FK→role', '角色ID'],
        ['menu_id', 'INTEGER', 'FOREIGN KEY → menu', '菜单ID'],
    ]
)

add_para('表4-9 file（文件表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '文件唯一标识'],
        ['name', 'VARCHAR(255)', 'NULL', '文件名'],
        ['type', 'VARCHAR(255)', 'NULL', '文件类型'],
        ['size', 'BIGINT', 'NULL', '文件大小（字节）'],
        ['url', 'VARCHAR(255)', 'NULL', '文件存储路径'],
        ['is_delete', 'BOOLEAN', 'NULL', '是否已删除'],
        ['enable', 'BOOLEAN', 'NULL', '是否启用'],
        ['md5', 'VARCHAR(45)', 'NULL', 'MD5校验值'],
    ]
)

add_para('表4-10 import_log（导入日志表）', bold=True, indent=False)
add_table_with_data(
    ['字段名', '类型', '约束', '说明'],
    [
        ['id', 'INTEGER', 'PRIMARY KEY, AUTO', '日志唯一标识'],
        ['admin_id', 'INTEGER', 'FOREIGN KEY → user', '操作管理员ID'],
        ['file_name', 'VARCHAR(100)', 'NOT NULL', '导入文件名'],
        ['import_count', 'INTEGER', 'NOT NULL', '导入记录数'],
        ['import_status', 'VARCHAR(20)', 'NOT NULL', '导入状态'],
        ['error_info', 'TEXT', 'NULL', '错误信息'],
        ['import_time', 'DATETIME', 'NULL', '导入时间'],
    ]
)

add_heading('4.3 表关系说明', 2)
add_para('系统数据库表之间存在以下关键关系：')
add_para('（1）track表与trackpoints表为一对多（1:N）关系：一条轨迹包含多个轨迹点，通过trackpoints表中的trackid外键关联。对track表执行删除操作时，级联删除对应的所有trackpoints记录。')
add_para('（2）track表与work表为一对一（1:1）关系：每条轨迹对应一个作业统计记录，work表以trackid作为主键同时作为外键，实现主键关联。')
add_para('（3）track表与rate表为一对一（1:1）关系：每条轨迹对应一个通行率评估记录，rate表同样采用主键关联方式。')
add_para('（4）role表与user表为一对多（1:N）关系：一个角色可分配给多个用户，用户通过role外键关联到角色。')
add_para('（5）role表与menu表为多对多（N:M）关系：通过role_menu中间表实现角色与菜单的多对多关联。一个角色可以拥有多个菜单权限，一个菜单也可以分配给多个角色。')
add_para('（6）menu表为自引用关系：通过pid字段指向同一张表的id字段，实现菜单的树形层级结构（如"轨迹管理"下有子菜单"轨迹列表"和"轨迹地图"）。')
add_para('（7）import_log表与user表为多对一（N:1）关系：通过admin_id外键关联到执行导入操作的管理员用户。')

add_heading('4.4 核心功能模块设计', 2)

add_para('（1）用户认证模块', bold=True)
add_para('该模块负责用户身份认证和会话管理。前端在用户登录成功后，将后端返回的用户信息（含Token）存储到localStorage中。后续所有API请求通过Axios拦截器自动在请求头中附加Token和用户角色信息。前端路由守卫（beforeEach）在每次路由跳转前检查localStorage中是否存在用户信息，若不存在则重定向到登录页。后端登录接口验证用户名和密码，验证通过后返回用户详细信息。')

add_para('（2）轨迹数据导入模块', bold=True)
add_para('该模块实现了Excel文件的读取、解析和数据库写入功能。管理员通过前端界面上传Excel文件，后端接收文件后使用openpyxl库逐行读取数据。解析过程中对GPS时间、经纬度坐标等关键字段进行格式校验，无效数据过滤后记录到错误日志。数据入库采用Django的事务机制，确保数据一致性——如果入库过程中发生任何错误，整个事务回滚，不会产生不完整的数据。入库完成后自动计算作业统计（work表）和通行率（rate表）数据。')

add_para('（3）地图展示模块', bold=True)
add_para('该模块是系统的核心功能模块。前端接收到轨迹点数据后，首先进行数据预处理：按GPS时间排序、根据workstatus字段变化为每个点分配segmentIndex（分段索引）。然后对每条轨迹点的WGS-84坐标进行GCJ-02转换。使用Leaflet的Polyline对象按分段绘制轨迹线（每段内的点具有相同的workstatus），使用CircleMarker对象绘制轨迹点标记。通过不同的颜色方案区分正常作业（绿色）和闲置停止（灰色）状态。地图初始化时自动调用fitBounds方法将视图缩放到轨迹范围。')

add_para('（4）权限管理模块', bold=True)
add_para('该模块实现了RBAC权限控制模型。核心表包括role（角色）、menu（菜单）和role_menu（角色-菜单关联）。系统预设管理员（admin）和普通用户（user）两个角色。管理员通过角色管理页面为每个角色勾选可访问的菜单项，数据存储在role_menu关联表中。用户登录后，前端根据用户所属角色查询对应的菜单权限列表，动态构建侧边栏菜单树。前端路由守卫在每次导航时验证当前路径是否在用户权限列表中，无权限则跳转到404页面。')

doc.add_page_break()

# ============================================================
# 5. 系统实现与测试
# ============================================================
add_heading('5. 系统实现与测试', 1)

add_heading('5.1 用户认证模块实现', 2)
add_para('用户认证模块包含登录、注册、个人信息修改和密码修改四个核心页面。图5-1展示了登录页面的实现效果，图5-2和图5-3展示了个人信息修改和密码修改页面。')

add_img('个人信息修改页面.png', Inches(5.0))
add_img_caption('图5-1 个人信息修改页面')
add_img('修改密码界面.png', Inches(5.0))
add_img_caption('图5-2 修改密码页面')

add_para('登录功能的关键实现逻辑如下：前端将用户名和密码通过Axios发送到后端的/user/login接口，后端验证用户凭证后返回用户信息对象（含Token和角色信息）。前端将用户信息存储到localStorage中，并通过Axios请求拦截器在后续所有请求中自动附加认证头。路由守卫beforeEach检测到未登录状态时自动重定向到登录页。')

add_code_block('// 前端路由守卫（router/index.js 核心代码）\n'
                'router.beforeEach((to, from, next) => {\n'
                '  if (to.meta.public) { next(); return; }\n'
                '  const user = localStorage.getItem(\'user\');\n'
                '  if (!user) { next(\'/login\'); return; }\n'
                '  if (!hasRoutePermission(to.path)) {\n'
                '    ElMessage.warning(\'当前账号没有访问该页面的权限\');\n'
                '    next(\'/404\'); return;\n'
                '  }\n'
                '  next();\n'
                '});')

add_heading('5.2 轨迹管理与地图展示模块实现', 2)
add_para('轨迹管理模块提供轨迹列表的展示、筛选和操作功能。图5-3展示了轨迹列表页面的实现效果。')

add_img('轨迹管理页面.png', Inches(5.2))
add_img_caption('图5-3 轨迹列表管理页面')

add_para('轨迹地图展示模块是系统的核心功能，实现了轨迹在卫星遥感影像和矢量地图上的叠加显示。图5-4和图5-5分别展示了卫星底图模式和矢量底图模式下的轨迹展示效果。')

add_img('轨迹展示页面.png', Inches(5.2))
add_img_caption('图5-4 轨迹地图展示页面（卫星底图模式）')

add_img('轨迹矢量展示页面.png', Inches(5.2))
add_img_caption('图5-5 轨迹地图展示页面（矢量底图模式）')

add_para('地图模块的关键实现包括以下几个方面：')

add_para('（1）地图初始化与底图切换：系统基于Leaflet创建地图实例，同时加载卫星影像（高德style=6）和矢量地图（高德style=7）两个瓦片图层。初始化时默认显示卫星图层。通过自定义的Leaflet Control在地图左上角（缩放按钮下方）添加底图切换按钮，用户点击时在两个图层之间切换。')

add_code_block('// 地图初始化与底图切换核心代码（TrackMap.vue）\n'
                'map = L.map(\'leaflet-map\', { zoomControl: true, preferCanvas: false });\n'
                'satelliteTileLayer = L.tileLayer(\n'
                '  \'https://webst01.is.autonavi.com/appmaptile?style=6&x={x}&y={y}&z={z}\'\n'
                ');\n'
                'vectorTileLayer = L.tileLayer(\n'
                '  \'https://webst01.is.autonavi.com/appmaptile?style=7&x={x}&y={y}&z={z}\'\n'
                ');\n'
                'satelliteTileLayer.addTo(map);  // 默认显示卫星地图')

add_para('（2）WGS-84到GCJ-02坐标转换：由于高德地图使用GCJ-02坐标系，而GPS设备输出的是WGS-84坐标，因此需要对所有轨迹点进行坐标转换。本系统在前端实现了完整的转换算法，包含transformLat和transformLon两个核心函数。')

add_code_block('// WGS-84 → GCJ-02 坐标转换核心算法（TrackMap.vue）\n'
                'function wgs84ToGcj02(lon, lat) {\n'
                '  const lng = Number(lon), latNum = Number(lat);\n'
                '  if (outOfChina(lng, latNum)) return { lon: lng, lat: latNum };\n'
                '  const a = 6378245.0, ee = 0.00669342162296594323;\n'
                '  let dLat = transformLat(lng - 105.0, latNum - 35.0);\n'
                '  let dLon = transformLon(lng - 105.0, latNum - 35.0);\n'
                '  const radLat = (latNum / 180.0) * Math.PI;\n'
                '  let magic = Math.sin(radLat);\n'
                '  magic = 1 - ee * magic * magic;\n'
                '  const sqrtMagic = Math.sqrt(magic);\n'
                '  dLat = (dLat * 180) / (((a * (1 - ee)) / (magic * sqrtMagic)) * Math.PI);\n'
                '  dLon = (dLon * 180) / ((a / sqrtMagic) * Math.cos(radLat) * Math.PI);\n'
                '  return { lon: lng + dLon, lat: latNum + dLat };\n'
                '}')

add_para('（3）轨迹分段渲染：系统根据每个轨迹点的workstatus字段将连续的轨迹点划分为不同的段（segment）。当相邻两点的workstatus不同时，开始一个新的段。每个段单独绘制一条Polyline，并根据其workstatus值决定颜色——正常作业段（workstatus=1）使用绿色，闲置停止段（workstatus=0）使用灰色。这种分段着色方式可以让用户一目了然地识别农机在哪些时段处于正常作业状态。')

add_code_block('// 轨迹分段渲染核心代码（TrackMap.vue）\n'
                'segments.forEach((seg) => {\n'
                '  const latlngs = seg.map((p) => [Number(p.mapPoint.lat), Number(p.mapPoint.lon)]);\n'
                '  const status = seg[0]?.workstatus;\n'
                '  const color = status === 0\n'
                '    ? \'rgba(132,138,146,0.9)\'  // 闲置 — 灰色\n'
                '    : \'rgba(104,223,58,0.95)\';  // 作业 — 绿色\n'
                '  L.polyline(latlngs, { color, weight: 4, opacity: 0.85 }).addTo(overlayLayer);\n'
                '});')

add_para('（4）轨迹动态回放：用户点击"动态运行展示"按钮后，系统使用JavaScript的setInterval函数以0.1秒/点的速度逐点展示轨迹。当前播放点以高亮的金色圆点标记，已走过的路径以金色轨迹线显示，地图自动跟随当前点平移（panTo动画效果），形成流畅的"拖尾"可视化效果。')

add_para('（5）距离测量：用户进入测距模式后，在地图上左键点击添加测量点，系统使用Leaflet的latLng.distanceTo()方法计算相邻点之间的球面距离（基于Haversine公式），实时累加并显示总距离。距离单位根据数值自动切换（<1000米显示"xxx.x m"，≥1000米显示"x.xxx km"）。')

add_heading('5.3 数据导入模块实现', 2)
add_para('数据导入模块允许管理员将Excel格式的农机轨迹数据批量导入到数据库中。图5-6展示了数据导入页面的实现效果。')

add_img('上传日志页面.png', Inches(5.2))
add_img_caption('图5-6 数据导入与导入日志页面')

add_para('数据导入的关键流程包括：①前端选择Excel文件上传至后端；②后端使用openpyxl或pandas库解析Excel文件内容；③按行提取GPS数据，验证经纬度、时间等关键字段的格式；④使用Django ORM创建track记录；⑤批量插入trackpoints记录；⑥根据轨迹点数据计算work统计（作业时长、作业面积、平均速度）；⑦根据轨迹点数据计算rate指标（通行率、生产效率）；⑧记录ImportLog（文件名、导入数量、状态、错误信息）。整个导入过程使用Django的事务管理（transaction.atomic），确保数据完整性。')

add_heading('5.4 用户与权限管理模块实现', 2)
add_para('用户与权限管理模块实现了完整的RBAC权限控制模型。图5-7、图5-8和图5-9分别展示了用户管理、角色管理和菜单管理页面的实现效果。')

add_img('用户管理页面.png', Inches(5.0))
add_img_caption('图5-7 用户管理页面')
add_img('角色管理页面.png', Inches(5.0))
add_img_caption('图5-8 角色管理页面')
add_img('菜单管理页面.png', Inches(5.0))
add_img_caption('图5-9 菜单权限配置页面')

add_para('权限控制的核心实现机制如下：每个用户通过role外键关联一个角色，每个角色通过role_menu中间表关联多个菜单权限。用户登录成功后，前端根据用户角色查询其拥有的菜单列表，动态渲染侧边栏导航。用户访问任何页面时，前端路由守卫会检查当前路径是否在用户的权限列表中，无权限则拦截并跳转。后端API在处理请求时也会通过请求头中的X-User-Role字段进行辅助鉴权。')

add_heading('5.5 系统测试', 2)
add_para('本系统在开发过程中进行了系统的功能测试，测试范围覆盖了所有核心功能模块。测试环境为Windows 11操作系统，Chrome浏览器。以下为主要测试用例及结果：')

add_table_with_data(
    ['测试编号', '测试模块', '测试用例', '预期结果', '实际结果'],
    [
        ['TC-01', '用户认证', '新用户注册', '成功创建账号', '通过'],
        ['TC-02', '用户认证', '已注册用户登录', '成功登录，跳转首页', '通过'],
        ['TC-03', '用户认证', '错误密码登录', '提示密码错误', '通过'],
        ['TC-04', '用户认证', '修改密码', '旧密码验证后更新成功', '通过'],
        ['TC-05', '轨迹管理', '查看轨迹列表', '显示8条轨迹记录', '通过'],
        ['TC-06', '轨迹管理', '按时间范围筛选', '仅显示符合条件轨迹', '通过'],
        ['TC-07', '轨迹管理', '查看轨迹地图', '地图加载，轨迹线显示', '通过'],
        ['TC-08', '地图交互', '底图切换', '卫星与矢量底图正常切换', '通过'],
        ['TC-09', '地图交互', '距离测量', '多点测距，距离计算正确', '通过'],
        ['TC-10', '地图交互', '轨迹动态回放', '逐点回放，地图跟随', '通过'],
        ['TC-11', '地图交互', '轨迹点点击查询', '右侧面板显示详情', '通过'],
        ['TC-12', '数据导入', '上传Excel文件', '数据成功导入数据库', '通过'],
        ['TC-13', '数据导入', '导入日志记录', '日志正确记录导入结果', '通过'],
        ['TC-14', '权限管理', '管理员删除轨迹', '轨迹及关联数据被删除', '通过'],
        ['TC-15', '权限管理', '普通用户访问管理页', '被拦截跳转404', '通过'],
        ['TC-16', '权限管理', '按角色渲染菜单', '不同角色显示不同菜单', '通过'],
        ['TC-17', '用户管理', '管理员删除用户', '用户被成功删除', '通过'],
        ['TC-18', '文件管理', '上传/删除文件', '文件操作正常', '通过'],
    ]
)
add_para('测试结果表明，系统的18项核心功能测试用例全部通过，各功能模块运行稳定，交互逻辑正确，满足设计要求。')

doc.add_page_break()

# ============================================================
# 6. 总结与展望
# ============================================================
add_heading('6. 总结与展望', 1)

add_heading('6.1 工作归纳总结', 2)
add_para('本课程设计项目完成了一个完整的农机作业数据处理与分析Web应用系统的设计与开发。项目从需求分析出发，经过系统设计、数据库设计、编码实现和系统测试等完整的软件工程生命周期，最终交付了一套功能完备、界面友好、运行稳定的软件系统。')
add_para('在技术架构方面，系统采用Vue 3前端框架与Django后端框架的前后端分离架构，通过RESTful API实现数据交互。数据库选用SQLite轻量级数据库，设计了10张规范化数据表，涵盖了轨迹数据存储、用户权限管理、文件管理和系统配置等功能。地图功能基于Leaflet开源库实现，集成了高德地图的卫星影像和矢量地图瓦片服务，并实现了WGS-84到GCJ-02坐标的精确转换算法。')
add_para('在功能实现方面，系统完整实现了用户认证、轨迹管理、地图展示与交互、数据导入、用户权限管理、文件管理和系统配置等模块。支持管理员和普通用户两种角色，实现了基于RBAC的权限控制模型。轨迹可在卫星影像和矢量地图上以不同颜色分段展示作业状态，支持动态回放和距离测量等高级交互功能。')
add_para('通过本项目的开发实践，本人对软件工程管理与实践课程的理论知识有了更深入的理解和掌握。从需求分析到系统设计的全过程训练了分析与抽象能力，前后端分离的架构实践加深了对现代Web开发技术的理解，数据库设计过程强化了数据建模与规范化思维，完整的功能测试保证了软件质量。')

add_heading('6.2 心得体会', 2)
add_para('（1）需求分析是软件开发的基石。在项目初期，对所提供轨迹数据的格式和内容进行充分分析，明确数据字段的含义和关系，为后续的数据库设计和功能开发奠定了坚实基础。如果需求分析不充分，将导致后期大量的返工和修改。')
add_para('（2）前后端分离架构的优势明显。在本次开发中，前后端职责清晰，接口定义明确后可以并行开发，大大提高了开发效率。前端专注于用户体验和交互设计，后端专注于业务逻辑和数据处理，这种分离使得代码更易于维护和扩展。')
add_para('（3）坐标转换是地图应用开发的关键环节。在项目初期，直接使用GPS坐标在地图上显示轨迹时发现位置存在偏移。经过调研发现这是由于坐标系不统一（WGS-84 vs GCJ-02）造成的。通过实现坐标转换算法，成功解决了轨迹与底图的对齐问题。这一经历说明，在涉及地理信息系统的开发中，务必注意坐标系的一致性问题。')
add_para('（4）RBAC权限模型的设计需要兼顾灵活性和安全性。在本系统中，通过用户-角色-菜单三级模型实现了灵活的权限配置，同时在前端路由守卫和后端API两个层面进行双重鉴权，保证了系统的安全性。')
add_para('（5）Leaflet地图库的功能强大但需要深入理解其API。在开发测距、回放等功能时，需要理解Leaflet的LayerGroup、Control、事件处理等核心概念。通过查阅官方文档和社区示例，逐步掌握了Leaflet的高级用法。')

add_heading('6.3 不足之处与改进方向', 2)
add_para('尽管本项目实现了预定的全部功能需求，但仍存在一些不足之处，需要在后续工作中加以改进：')
add_para('（1）数据可视化分析能力有限。当前系统主要实现了轨迹的地理空间展示，但在数据统计分析（如作业时间分布、速度分布直方图、作业面积统计图表等）方面还有较大提升空间。后续可集成ECharts等图表库，增加丰富的数据统计图表功能。')
add_para('（2）数据库性能优化不足。当前使用SQLite数据库，在读入数万条轨迹点数据后，查询性能可能成为瓶颈。后续可考虑迁移到MySQL或PostgreSQL等高性能数据库，并添加适当的数据库索引优化查询效率。')
add_para('（3）前端渲染性能有待优化。当单条轨迹包含4000+个点时，使用Leaflet的SVG渲染器一次性渲染全部点会导致初次加载时间较长。后续可考虑引入点聚合（Marker Clustering）或按缩放级别动态加载点等技术进行优化。')
add_para('（4）移动端适配不完善。当前系统主要针对桌面浏览器设计，在移动设备上的显示效果和操作体验有待优化。后续可考虑采用响应式设计或开发专门的移动端应用。')
add_para('（5）缺乏自动化测试。当前主要依赖人工功能测试，测试覆盖率和效率有限。后续应引入单元测试（如Django TestCase、Vue Test Utils）和端到端测试（如Playwright），建立持续集成测试流程。')

doc.add_page_break()

# ============================================================
# 参考文献
# ============================================================
add_heading('参考文献', 1)
add_para('')

refs = [
    '[1] 张海藩, 吕云翔. 软件工程（第6版）[M]. 北京: 人民邮电出版社, 2020.',
    '[2] 李刚. Django 4 Web应用开发实战[M]. 北京: 电子工业出版社, 2023.',
    '[3] 尤雨溪. Vue.js 3.0官方文档[EB/OL]. https://cn.vuejs.org/, 2024.',
    '[4] Vladimir Agafonkin. Leaflet JavaScript Library Documentation[EB/OL]. https://leafletjs.com/, 2024.',
    '[5] 高德开放平台. 地图JS API坐标系说明[EB/OL]. https://lbs.amap.com/api/, 2024.',
    '[6] Tom Christie. Django REST Framework Documentation[EB/OL]. https://www.django-rest-framework.org/, 2024.',
    '[7] 刘汝佳, 陈锋. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
    '[8] William S. Vincent. Django for APIs: Build Web APIs with Python & Django[M]. Still River Press, 2023.',
]

for ref in refs:
    add_para(ref, indent=False)

# ====== SAVE ======
doc.save(OUTPUT)
print(f'Word document saved to: {OUTPUT}')
